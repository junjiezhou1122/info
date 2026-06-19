from pathlib import Path
import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path(
    "/Users/junjie/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_16qi1otyl6v921_bb36/msg/file/2026-06/"
    "MetaFlow无问芯穹赛题正式技术文档(1).docx"
)
OUT = Path("/Users/junjie/info/MetaFlow无问芯穹赛题正式技术文档-DeepResearch增强版.docx")
INDEX = Path("/Users/junjie/info/MetaFlow-DeepResearch-引用与竞品索引.md")


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "宋体"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(p, size=10.5, first_line=True, space_after=6):
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(space_after)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    for run in p.runs:
        set_run_font(run, size=size)


def style_heading(p, level=2):
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True, color="1F4E79")
    p.paragraph_format.space_before = Pt(10 if level == 2 else 14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = None


def insert_after(anchor, new_element):
    anchor._p.addnext(new_element._p)
    return new_element


def add_paragraph_after(anchor, text="", style=None, size=10.5, first_line=True):
    new_p = anchor.insert_paragraph_before(text, style=style)
    anchor._p.addnext(new_p._p)
    style_paragraph(new_p, size=size, first_line=first_line)
    return new_p


def add_heading_after(anchor, text, level=2):
    style = f"Heading {level}"
    p = add_paragraph_after(anchor, text, style=style, size=12, first_line=False)
    style_heading(p, level=level)
    return p


def add_heading_before(anchor, text, level=2):
    style = f"Heading {level}"
    p = anchor.insert_paragraph_before(text, style=style)
    style_heading(p, level=level)
    return p


def set_cell_text(cell, text, bold=False, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table_after(anchor, headers, rows, widths=None):
    doc = anchor.part.document
    table = doc.add_table(rows=1, cols=len(headers))
    for style_name in ("Table Grid", "网格型"):
        try:
            table.style = style_name
            break
        except KeyError:
            pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=8.5)
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=8)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    anchor._p.addnext(table._tbl)
    return table


def add_paragraph_after_table(table, text="", style=None, size=10.5, first_line=True):
    doc = table.part.document
    p = doc.add_paragraph(text, style=style)
    table._tbl.addnext(p._p)
    style_paragraph(p, size=size, first_line=first_line)
    return p


def find_paragraph(doc, exact_text):
    for p in doc.paragraphs:
        if p.text.strip() == exact_text:
            return p
    raise ValueError(f"paragraph not found: {exact_text}")


def find_heading_startswith(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(f"heading not found: {prefix}")


def replace_paragraph_text(doc, old, new, size=10.5):
    p = find_paragraph(doc, old)
    p.text = new
    style_paragraph(p, size=size)
    return p


def rename_heading(doc, old, new):
    p = find_paragraph(doc, old)
    p.text = new
    style_heading(p, level=2)


def add_reference_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    for run in p.runs:
        set_run_font(run, size=9)


aha_heading = "1.4 核心洞察：从 Agent Demo 到上下文运行时"
aha_paragraphs = [
    "MetaFlow 的 aha moment 不在于“又做了一个 Agent”，而在于把 Agent 产业正在缺失的一层补出来：把脚本、浏览器、CLI、外部 Agent、个人应用和长期上下文统一成可组合、可验证、可跨端运行的 runtime/adapter layer。Demo 时代比的是 Agent 会不会做事；runtime 时代比的是新 Agent、新工具、新端能不能不重写系统地接进来。MCP 将工具与数据源抽象为可发现、可调用的外部能力，ACP 将客户端与 Agent 运行过程解耦，A2A 则进一步强调不同 Agent 应用之间的互操作边界[5-7]。MetaFlow 顺着这一行业趋势，把端侧事实、任务视图、执行器、外部工具和云端模型放进同一套证据与权限模型。",
    "MetaFlow 的基本单位不是 message，而是 View。message 只是发生过的对话或事件，View 是下一次工作能直接站上去的状态。许多记忆系统把历史材料塞进向量库，希望未来检索“刚好命中”；MetaFlow 则把当前页面、项目状态、失败证据、用户反馈、下一步行动物化为可检查、可分叉、可归档、可复用的 ViewGraph。LangGraph 等框架已经把持久化、human-in-the-loop 和 long-running agent 作为核心能力[8]，Anthropic 也强调上下文工程需要主动设计而非简单堆入提示词[9]。MetaFlow 将这些原则前移到端侧/云端协同的上下文底座中。",
    "不可验证的 Agent 只能演示，可验证的 Agent 才能成为运行时。单点 Agent 最大的问题不是不能成功，而是失败时不知道为什么，成功时也难以复用。OpenAI Agents SDK、OpenTelemetry GenAI 语义约定和 ACP 工具调用协议都把 tracing、tool call lifecycle、权限请求、状态更新推向核心位置[10-12]。MetaFlow 将每次观察、证据选择、模型调用、工具执行、用户拒绝和人工修改写回 ViewGraph，使其既是审计日志，也是后续学习和工作流沉淀的证据。",
    "因此，MetaFlow 选择“脚本确定性 + Agent 探索性 + ViewGraph 记忆 + feedback 演化”的中间路线。Anthropic 将 workflow 与 agent 区分为两类工程模式：可预测任务优先用 workflow，开放任务再交给 agent 探索[13]。MetaFlow 把二者放入同一个运行框架：重复成功的路径沉淀为 processor 或工作流，失败路径保留原因并回到 Agent 探索，用户反馈持续修正上下文选择、权限边界和自动化程度。",
]

related_heading = "1.6 国内外相关工作与差异化定位"
related_intro = [
    "当前 AI Agent 生态可以分为三层：第一层是浏览器/桌面执行器，例如 browser-use、Skyvern、Browserbase Stagehand、trycua/cua、OpenCUA 等，解决“Agent 如何使用浏览器或电脑”的动作空间问题[14-19]；第二层是 Agent 编排框架，例如 LangGraph、CrewAI、AutoGen、AutoGPT、OpenHands，解决多步骤规划、工具调用、长运行状态和开发任务执行问题[8,20-23]；第三层是评测与数据集，例如 WebArena、Mind2Web、OSWorld、AndroidWorld、GAIA、SWE-bench 等，用可复现任务衡量 Agent 是否真的完成了目标[24-30]。这些项目共同说明：Agent 能力正在从一次性聊天走向真实工具环境，但可靠性、可控性和可复用性仍是工程瓶颈。",
    "从商业产品看，OpenAI Operator/Computer-Using Agent、Anthropic Claude Computer Use、Google Project Mariner、Manus、Devin、Replit Agent、Browserbase、Firecrawl、browserless、MultiOn 等产品已经证明 browser-use、computer-use 和 autonomous task execution 是明确方向[31-40]。但多数方案要么是闭源托管黑盒，要么只提供浏览器基础设施，要么锁定 coding、办公或网页自动化的单一场景。MetaFlow 的差异化在于把这些能力收束为个人上下文智能运行框架：端侧保留事实和权限，云端完成复杂理解，外部执行器通过 adapter 接入，工作流通过证据与反馈持续验证。",
]

open_source_rows = [
    ("OpenClaw", "自托管个人 AI 助手，面向聊天入口、技能和本机动作。", "个人助理生态强；MetaFlow 更强调 ViewGraph、证据包、权限审计和工作流沉淀。", "[14]"),
    ("browser-use", "让 Agent 使用真实浏览器的开源浏览器自动化库。", "可作为浏览器执行器；MetaFlow 负责跨任务状态、证据、反馈与端云协同。", "[15]"),
    ("Skyvern", "AI-powered browser workflow automation，适合表单和多步网页流程。", "偏浏览器 RPA；MetaFlow 将浏览器、CLI、文件、模型、人工确认统一编排。", "[16]"),
    ("Browserbase Stagehand", "自然语言 + Playwright 风格代码的浏览器 Agent SDK。", "SDK 层能力强；MetaFlow 补运行时层的状态、权限、日志、回放和复用。", "[17]"),
    ("trycua/cua / OpenCUA", "面向 computer-use agent 的 sandbox、SDK、数据和评测。", "可作为桌面执行或评测后端；MetaFlow 提供上层任务证据包和运行策略。", "[18][19]"),
    ("OpenHands", "开源软件开发 Agent 平台。", "聚焦代码任务；MetaFlow 面向更广泛的研究、写作、浏览器、项目续接场景。", "[22]"),
    ("LangGraph", "长运行、有状态 Agent 编排框架。", "开发框架能力强；MetaFlow 在其上层强调端侧事实、个人上下文和竞赛/产品闭环。", "[8]"),
    ("CrewAI / AutoGen / AutoGPT", "多 Agent 自动化与自主 Agent 代表框架。", "强调 Agent 协作；MetaFlow 进一步要求可验证、可回退、可沉淀的 workflow runtime。", "[20][21][23]"),
]

commercial_rows = [
    ("OpenAI Operator / CUA", "基于视觉与强化学习推理的浏览器/GUI 行动 Agent。", "模型与托管产品强；MetaFlow 强调自有上下文、可替换模型和本地/云端可审计执行。", "[31]"),
    ("Claude Computer Use", "通过截图、鼠标、键盘控制桌面环境的模型级 API。", "能力仍需沙箱、权限、审计和 agent loop；MetaFlow 将这些变成系统结构。", "[32]"),
    ("Google Project Mariner", "Gemini 浏览器 Agent 原型，可理解网页像素、文本、代码和表单。", "研究原型与 Chrome 扩展路径；MetaFlow 将浏览器插件能力协议化、可复现化。", "[33]"),
    ("Browserbase / browserless", "云端浏览器基础设施，支持会话、代理、反爬和 Playwright/Puppeteer。", "解决浏览器运行；MetaFlow 解决任务生命周期、证据链、权限和反馈闭环。", "[34][35]"),
    ("Firecrawl", "Search、scrape、crawl、interact 的 web context API。", "可作为资料获取 adapter；MetaFlow 将网页证据纳入 ViewGraph 和模型证据包。", "[36]"),
    ("MultiOn / Manus", "托管 autonomous web/general AI agent API。", "交付方便但闭源黑盒；MetaFlow 更适合竞赛复现、私有化和可控工作流。", "[37][38]"),
    ("Devin / Replit Agent / Google Jules", "工程任务或应用生成 Agent。", "垂直 coding 能力强；MetaFlow 可把 coding agent 作为一类 processor 接入。", "[39][40][41]"),
]

benchmark_paragraphs = [
    "公开基准也解释了为什么 MetaFlow 必须强调“可验证运行时”而不是单次 Agent 演示。WebArena 在自托管真实网站环境中报告 GPT-4 基线端到端成功率明显低于人类，说明真实网页任务的长程规划、状态验证和错误恢复仍是瓶颈[24]。Mind2Web 覆盖 137 个网站、31 个领域，强调跨网站泛化；Online-Mind2Web 进一步指出静态评测会高估 Agent 能力，因为真实网页存在弹窗、cookie、布局变化和在线状态漂移[25][26]。OSWorld 与 AndroidWorld 则把问题扩展到桌面与移动端，要求 Agent 处理 GUI grounding、应用状态、触控和跨应用任务[28][29]。",
    "这些研究给 MetaFlow 的测试方案带来三个约束。第一，评估不能只看最终回答，要记录观察、工具调用、失败原因、人工接管和状态回放。第二，必须把 Web、桌面、移动、API、文件和代码任务纳入统一证据模型，否则不同执行器之间无法比较和复用经验。第三，商业化场景要采用渐进自治：先以建议、草稿、只读抽取和低风险自动化建立可信闭环，再把经过验证的重复路径沉淀为工作流。GAIA、AgentBench、SWE-bench、WorkArena、ToolSandbox 等基准都表明，真实任务需要推理、工具、状态、约束和人类反馈共同闭环[27,30,42-45]。",
]

references = [
    "[1] 无问芯穹. 无问芯穹平台与大模型基础设施相关资料[EB/OL].",
    "[2] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. NeurIPS, 2017. https://arxiv.org/abs/1706.03762",
    "[3] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]. ICLR, 2023. https://arxiv.org/abs/2210.03629",
    "[4] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020. https://arxiv.org/abs/2005.11401",
    "[5] Model Context Protocol. Introduction[EB/OL]. https://modelcontextprotocol.io/docs/getting-started/intro",
    "[6] Agent Client Protocol. Introduction and Architecture[EB/OL]. https://agentclientprotocol.com/get-started/introduction",
    "[7] A2A Project. Agent2Agent Protocol[EB/OL]. https://github.com/a2aproject/A2A",
    "[8] LangChain. LangGraph Overview and Persistence[EB/OL]. https://docs.langchain.com/oss/python/langgraph/overview",
    "[9] Anthropic. Effective context engineering for AI agents[EB/OL]. https://www.anthropic.com/engineering/effective-context-engineering-for-ai-agents",
    "[10] OpenAI. Agents SDK tracing[EB/OL]. https://openai.github.io/openai-agents-python/tracing/",
    "[11] OpenTelemetry. GenAI semantic conventions[EB/OL]. https://opentelemetry.io/docs/specs/semconv/registry/attributes/gen-ai/",
    "[12] Agent Client Protocol. Tool calls and session updates[EB/OL]. https://agentclientprotocol.com/protocol/tool-calls",
    "[13] Anthropic. Building effective agents[EB/OL]. https://www.anthropic.com/engineering/building-effective-agents",
    "[14] OpenClaw. OpenClaw GitHub repository[EB/OL]. https://github.com/openclaw/openclaw",
    "[15] browser-use. browser-use GitHub repository[EB/OL]. https://github.com/browser-use/browser-use",
    "[16] Skyvern-AI. Skyvern GitHub repository[EB/OL]. https://github.com/Skyvern-AI/skyvern",
    "[17] Browserbase. Stagehand GitHub repository[EB/OL]. https://github.com/browserbase/stagehand",
    "[18] trycua. CUA GitHub repository[EB/OL]. https://github.com/trycua/cua",
    "[19] XLang Lab. OpenCUA[EB/OL]. https://github.com/xlang-ai/OpenCUA",
    "[20] CrewAI. CrewAI GitHub repository[EB/OL]. https://github.com/crewAIInc/crewAI",
    "[21] Wu Q, Bansal G, Zhang J, et al. AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation Framework[EB/OL]. https://arxiv.org/abs/2308.08155",
    "[22] OpenHands. OpenHands GitHub repository[EB/OL]. https://github.com/OpenHands/openhands",
    "[23] Significant Gravitas. AutoGPT Classic README[EB/OL]. https://github.com/Significant-Gravitas/AutoGPT",
    "[24] Zhou S, Xu F F, Zhu H, et al. WebArena: A Realistic Web Environment for Building Autonomous Agents[EB/OL]. https://arxiv.org/abs/2307.13854",
    "[25] Deng X, Gu Y, Zheng B, et al. Mind2Web: Towards a Generalist Agent for the Web[EB/OL]. https://arxiv.org/abs/2306.06070",
    "[26] Online-Mind2Web. An Illusion of Progress? Assessing the Current State of Web Agents[EB/OL]. https://arxiv.org/html/2504.01382v4",
    "[27] Mialon G, Fourrier C, Swift C, et al. GAIA: A Benchmark for General AI Assistants[EB/OL]. https://arxiv.org/abs/2311.12983",
    "[28] XLang Lab. OSWorld: Benchmarking Multimodal Agents for Open-Ended Tasks in Real Computer Environments[EB/OL]. https://arxiv.org/abs/2404.07972",
    "[29] Google Research. AndroidWorld: A Dynamic Benchmarking Environment for Autonomous Agents[EB/OL]. https://arxiv.org/abs/2405.14573",
    "[30] Liu X, Yu H, Zhang H, et al. AgentBench: Evaluating LLMs as Agents[EB/OL]. https://arxiv.org/abs/2308.03688",
    "[31] OpenAI. Computer-Using Agent and Operator[EB/OL]. https://openai.com/index/computer-using-agent/",
    "[32] Anthropic. Claude computer use tool documentation[EB/OL]. https://platform.claude.com/docs/en/agents-and-tools/tool-use/computer-use-tool",
    "[33] Google DeepMind. Project Mariner and Gemini 2.0 update[EB/OL]. https://blog.google/innovation-and-ai/models-and-research/google-deepmind/google-gemini-ai-update-december-2024/",
    "[34] Browserbase. Browserbase official website[EB/OL]. https://www.browserbase.com/",
    "[35] browserless. browserless official website[EB/OL]. https://www.browserless.io/",
    "[36] Firecrawl. Firecrawl official website and GitHub repository[EB/OL]. https://www.firecrawl.dev/",
    "[37] MultiOn. MultiOn documentation[EB/OL]. https://docs.multion.ai/welcome",
    "[38] Manus. Manus documentation and API[EB/OL]. https://manus.im/docs/introduction/welcome",
    "[39] Cognition. Introducing Devin[EB/OL]. https://www.cognition.ai/blog/introducing-devin",
    "[40] Replit. Replit Agent[EB/OL]. https://replit.com/ai",
    "[41] Google. Jules official website[EB/OL]. https://jules.google/",
    "[42] Jimenez C E, Yang J, Wettig A, et al. SWE-bench: Can Language Models Resolve Real-World GitHub Issues?[EB/OL]. https://arxiv.org/abs/2310.06770",
    "[43] ServiceNow Research. WorkArena and WorkArena++[EB/OL]. https://github.com/ServiceNow/workarena",
    "[44] Apple. ToolSandbox: A Stateful, Conversational, Interactive Evaluation Benchmark for LLM Tool Use Capabilities[EB/OL]. https://arxiv.org/abs/2408.04682",
    "[45] Sierra. tau-bench: A Benchmark for Tool-Agent-User Interaction in Real-World Domains[EB/OL]. https://arxiv.org/abs/2406.12045",
]

replacements = {
    "论文题目：MetaFlow：端侧/云端协同的个人上下文智能运行框架":
        "论文题目：MetaFlow：端云协同的个人 AI Agent 上下文运行时",
    "MetaFlow: An Edge-Cloud Collaborative Personal Context Intelligence Framework":
        "MetaFlow: A Local-First Edge-Cloud Context Runtime for Personal AI Agents",
    "本文面向无问芯穹端侧/云端协同智能应用赛道，提出 MetaFlow 个人上下文智能运行框架。MetaFlow 围绕用户真实工作过程构建“静默观察—理解处理—内部视图—智能行动—反馈学习”闭环系统。系统在端侧持续捕获屏幕活动、输入操作、文件变化、应用与终端事件等事实证据，将其整理为可追溯的观察结果，再压缩为面向任务的内部视图，使云端大模型和智能代理能够在受控、可解释、可复用的上下文基础上完成研究、写作、项目续接和知识治理等复杂任务。":
        "本文面向无问芯穹端侧/云端协同智能应用赛道，提出 MetaFlow：一个本地优先、端云协同的个人 AI Agent 上下文运行时。本文的核心判断是：个人智能应用的瓶颈不只是模型是否更聪明，而是用户长期工作状态能否从零散记录变成有类型、可检查、可复用、可验证的上下文数据产品。MetaFlow 将对话、浏览器活动、项目文件、日志、屏幕证据和用户反馈等原始工作痕迹规范为 Observation，再由确定性规则、脚本、LLM 与 agent-task processor 生成 state.surface、work.focus_set、project.current、task.*、result.*、feedback.*、memory.* 等任务 Views，形成“Observation -> Processor -> ViewGraph -> 应用/Agent 行动 -> Feedback -> Evolution”的持续闭环。",
    "MetaFlow 的主要创新在于把个人智能应用从单次问答提升为持续运行的上下文智能底座。端侧负责低延迟观察、隐私边界控制和本地证据沉淀；无问芯穹平台侧负责高复杂度模型推理、长上下文理解、多模态处理、智能代理规划和算力调度。系统通过“事实保真、视图生成、证据溯源、权限控制、反馈学习”的闭环机制，在保证隐私和可信性的前提下降低重复整理成本和模型输入成本，为研究辅助、项目续接、写作增强、个人知识治理和团队上下文中台等场景提供可扩展基础。":
        "与传统 RAG、笔记软件或单次 Agent demo 不同，MetaFlow 追求的不是“存下更多历史”，而是让下一次任务更便宜：更少搜索、更少重复解释、更少 token、更少失败重来。系统在端侧保留事实、来源和权限边界，在无问芯穹平台侧完成长上下文、多模态理解、复杂推理和智能代理探索；Agent 的输出不会静默成为事实，而是带着来源、运行轨迹和用户反馈写回 ViewGraph。由此，MetaFlow 将脚本的确定性、Agent 的探索性、ViewGraph 的长期状态和反馈驱动的工作流演化连接起来，为研究辅助、项目续接、写作增强、个人知识治理和团队上下文中台提供可落地的基础框架。",
    "关键词：MetaFlow；端侧智能；云端协同；无问芯穹平台；上下文智能；大模型；智能代理":
        "关键词：MetaFlow；端侧智能；云端协同；无问芯穹平台；ViewGraph；上下文运行时；智能代理；工作流沉淀",
    "This paper presents MetaFlow, an edge-cloud collaborative personal context intelligence framework for the Infini-AI track. MetaFlow builds a continuous loop of silent observation, structured understanding, internal views, intelligent action, and feedback learning. It preserves factual evidence on the edge, organizes task-oriented context with provenance, and invokes cloud models and agents through scoped evidence packages. With Infini-AI platform resources, MetaFlow supports research, writing, project continuation, and knowledge governance while improving context reuse, controllability, and model invocation efficiency.":
        "This paper presents MetaFlow, a local-first and edge-cloud collaborative personal AI agent context runtime for the Infini-AI track. Its core insight is that the bottleneck of personal AI is not only model intelligence, but whether long-running work traces can become typed, inspectable, reusable, and verifiable context products. MetaFlow converts conversations, browser activity, project files, logs, screen evidence, and user feedback into Observations, then uses deterministic, script, LLM, and agent-task processors to materialize task-specific Views such as state.surface, work.focus_set, project.current, task.*, result.*, feedback.*, and memory.*. These Views form an Observation -> Processor -> ViewGraph -> Action -> Feedback -> Evolution loop. With Infini-AI platform resources, MetaFlow keeps factual evidence and privacy boundaries on the edge while using cloud models for long-context reasoning, multimodal understanding, workflow exploration, and evidence-grounded generation.",
    "Keywords: edge AI; cloud collaboration; AI cloud service; large model; agent":
        "Keywords: edge AI; cloud collaboration; context runtime; ViewGraph; AI agent; workflow crystallization",
    "随着大语言模型、多模态感知和智能代理技术快速发展，端侧/云端协同已经成为智能应用的重要演进方向。传统人工智能应用往往以单次输入和单次输出为基本形态，模型能够处理当前问题，却难以理解用户在较长时间尺度上的任务背景、资料来源、工作状态和偏好变化。对于开发者、研究者、写作者、产品经理等知识工作者而言，真正影响任务质量的信息通常分散在网页阅读、文档编辑、会议记录、项目资料、历史对话和多轮修改过程中。如果这些上下文不能被持续保留和有效组织，智能代理很容易重复询问、遗漏证据、生成无来源结论，甚至在不同任务之间产生冲突。":
        "大语言模型、多模态感知和智能代理正在把“回答问题”推进到“使用工具完成任务”。OpenAI Operator/Computer-Using Agent、Anthropic Claude Computer Use、Google Project Mariner 等产品说明，模型已经开始通过浏览器和图形界面行动[31-33]；WebArena、Mind2Web、OSWorld、AndroidWorld 等基准则说明，真实网页、桌面和移动任务仍然存在长程规划、状态漂移、GUI grounding、错误恢复和人工接管难题[24-29]。这意味着个人 AI 的下一步突破不只是让模型更会点击，而是要给 Agent 一个可持续运行的上下文底座。",
    "端侧设备天然贴近用户真实工作现场，能够低延迟地获得浏览、编辑、阅读和交互过程中的事实证据；云端平台则具备更强的大模型推理、多智能代理协同、批量计算和长上下文处理能力。MetaFlow 将二者组织为持续运行的上下文智能闭环：端侧负责事实保真、隐私边界和即时交互，云端负责复杂理解、任务规划和工作流沉淀，用户反馈再反向影响后续上下文选择、模型调用和自动化程度。":
        "知识工作中的关键上下文并不天然存在于某个聊天窗口里，而是散落在浏览器页面、文档草稿、终端日志、项目文件、会议记录、历史对话、截图和多轮修改之中。普通 Agent 每次任务都像重新开局：用户重新解释背景，模型重新拼接材料，失败经验难以复用。MetaFlow 的出发点是把这些工作痕迹先当作事实证据，而不是马上当作“记忆”；端侧负责低延迟观察、隐私边界和事实保真，云端负责复杂理解、多模态处理、智能代理探索和工作流压缩，用户反馈再决定哪些 View、processor 和 workflow 值得保留或演化。",
    "本文研究的问题可概括为：如何在保护隐私和保证可解释性的前提下，将用户连续工作过程中产生的多源信息转化为可供大模型与智能代理复用的任务上下文。原始记录需要经过来源标注、权限约束、语义压缩和任务筛选才能成为可用知识；不同任务也需要不同粒度、不同来源和不同权限范围的证据。浏览器、写作、研究、项目续接和每日总结等场景可以共享同一批事实证据，由此形成统一的上下文智能基础。":
        "本文研究的问题可概括为：如何在保护隐私和保证可解释性的前提下，将连续工作过程中产生的多源事实转化为未来任务能直接使用的 typed Views。这里的目标不是保存尽可能多的历史，而是降低未来搜索成本和任务失败率。一个 Observation 可以被多个 processor 解释成不同 Views；同一批事实可以服务研究证据、项目续接、写作草稿、浏览器任务、每日记忆和团队知识沉淀；一个 View 只有在后续任务中减少搜索、减少 token、减少重复失败或提升完成质量时，才值得被保留、合并、提升或沉淀为 memory。",
    "MetaFlow 将该问题拆解为四个核心子问题。第一，如何区分事实与推断，避免在采集阶段把不稳定解释固化为长期记忆。第二，如何把高频、异构、碎片化的端侧信息压缩成面向任务的内部视图。第三，如何在调用云端大模型和智能代理时控制可见范围、模型输入成本和证据来源。第四，如何把用户的接受、修改、忽略和纠错转化为系统未来行为的学习信号。":
        "MetaFlow 将该问题拆解为四个核心子问题。第一，如何区分 Observation 与 View，避免在采集阶段把不稳定推断写成长期事实。第二，如何通过 ViewSpec 和 ProcessorSpec 把当前屏幕、工作焦点、项目状态、任务结果、反馈和记忆表达为可检查、可分叉、可归档、可追溯的 ViewGraph。第三，如何通过任务证据包和 agent runtime adapter 控制云端模型、浏览器执行器、CLI、MCP/ACP 工具的可见范围、权限和成本。第四，如何把接受、修改、拒绝、失败、复用和人工接管等反馈转化为 View 排名、processor 路由和 workflow crystallization 的演化信号。",
    "MetaFlow 定位为面向个人与团队知识工作的上下文智能基础框架。系统以“事实保真、内部视图、任务证据包、智能代理协同、反馈学习”为主线，将用户工作过程中的网页、文档、屏幕活动、项目资料、对话记录和操作反馈组织为可追溯的上下文网络，并按任务需求生成受控输入，交由无问芯穹平台的大模型推理、智能代理开发和算力服务处理。":
        "MetaFlow 定位为面向个人与团队知识工作的本地优先上下文运行时，而不是单一聊天助手、笔记软件或浏览器自动化工具。系统的基本对象是 View：当前状态是 View，任务是 View，结果是 View，反馈是一组 Views，memory 是一种会改变未来行为的 retained View。应用不是各自拥有记忆的孤岛，而是同一个 ViewGraph 的不同投影；Agent 也不是系统本身，而是可通过 CLI、HTTP、MCP、ACP 或 agent runtime adapter 接入的能力与外部 worker。",
    "在应用层面，MetaFlow 可支持研究资料整理、项目上下文续接、写作辅助、个人知识治理、每日工作摘要、团队知识沉淀等场景。在技术层面，系统强调端侧与云端各司其职：端侧保留事实与权限，云端完成理解与生成；端侧提供连续上下文，云端提供复杂推理；端侧保障隐私和低延迟，云端提供规模化智能能力。":
        "在应用层面，MetaFlow 可支持研究资料整理、项目上下文续接、写作辅助、个人知识治理、每日工作摘要、团队知识沉淀、浏览器任务 cockpit、workflow miner 等场景。每个应用都遵循“ViewSpec + Processor + App Surface + Feedback Contract”的模式：读取同一套 ViewGraph，渲染不同任务界面，写回可学习的反馈。这样，系统不会为每个应用复制一套记忆，而是让同一份上下文在不同端、不同任务、不同 Agent 之间被安全复用。",
    "第一，提出面向个人工作流的上下文智能框架。MetaFlow 将用户真实工作过程中不断产生的事实证据组织为可复用上下文，使大模型能够基于持续积累的任务背景开展推理。该思路将智能应用的重点从“模型回答能力”扩展到“上下文生产、压缩、治理和复用能力”。":
        "第一，提出 View-first 的个人 AI Agent 上下文运行时。MetaFlow 将智能应用的基本对象从 message、prompt 或单次任务提升为可物化的 View：当前屏幕、工作焦点、项目状态、任务结果、用户反馈和长期记忆都可以成为有类型、带来源、可查询、可分叉、可归档的上下文数据产品。该思路把智能应用的重点从“模型回答能力”扩展到“上下文生产、压缩、治理、复用和演化能力”。",
    "第二，提出事实层与语义层分离的上下文建模方法。系统将端侧捕获到的网页、文本、屏幕、项目和交互行为视为事实层，将摘要、意图、任务、工作流和长期偏好视为语义层。事实层尽量保持稳定和可追溯，语义层允许随着新证据和用户反馈不断更新，从而避免原始证据被模型推断污染。":
        "第二，提出 Observation 与 View 分离的可追溯建模方法。Observation 只记录发生了什么、来自哪里、何时发生、如何采集，不在采集阶段固化不稳定解释；View 则是面向任务的压缩、解释、组织或应用就绪状态。这样，同一条事实可以被不同 processor 重新解释为研究证据、项目状态、写作材料、失败模式或长期偏好，错误 View 也可以被归档、重建或替换，而不会污染原始事实。",
    "第三，提出工作流沉淀机制。对于未知或复杂任务，系统允许云端智能代理进行探索，记录成功路径和失败原因；当某类任务被多次验证后，将其沉淀为可复用的工作流方案。后续同类任务优先执行已验证流程，失败时再回退到智能代理探索。该机制在灵活性和稳定性之间取得平衡，降低重复模型输入消耗和人工干预成本。":
        "第三，提出 Agent Runtime Adapter 与协议化能力接入机制。MetaFlow 不把某一个 Agent 当作系统本身，而是把 Agent 视为可替换的外部 worker 或 capability。通过 CLI、HTTP、MCP、ACP、浏览器插件、Firecrawl 等 adapter，系统可以把浏览器、桌面、网页抓取、代码检查和外部 Agent runtime 接入同一套任务证据包、权限策略、运行事件和结果 View 中，使新工具、新模型和新端不需要重写核心上下文系统。",
    "第四，提出基于证据包的端云协同调用模式。云端模型接收经过目标筛选、权限过滤、来源标注和模型输入预算控制的任务证据包，在有限可见范围内完成推理。该模式既能保证云端智能能力获得足够背景，又能显著减少无关信息和敏感信息外发，提高推理可信度和系统可控性。":
        "第四，提出 workflow crystallization 与可验证反馈闭环。未知任务先允许 Agent 探索，并记录工具调用、失败原因、输出结果和用户反馈；当某类路径反复成功后，系统将其沉淀为可复用 processor 或 workflow recipe；后续同类任务优先执行确定性流程，失败时再回退给 Agent 探索。由此，脚本提供确定性，Agent 提供探索性，ViewGraph 提供长期状态，feedback 提供演化方向，系统会在真实使用中逐步降低 token、延迟、搜索步骤和重复失败。",
    "MetaFlow 的总体框架以“持续观察、结构理解、按任务行动、反馈再学习”为核心思想。图2-1展示了系统从静默观察源到结构化观察结果，再到理解处理、内部视图、行动渠道和新观察反馈的完整闭环，说明 MetaFlow 如何把分散的数字行为转化为可被人和智能代理共同使用的上下文能力。":
        "MetaFlow 的总体架构以“让上下文可检查，让记忆可行动，让反馈推动系统变聪明”为核心思想。图2-1展示了系统从 Observation stream 到 Processor runtime、ViewGraph、应用/Agent 行动、反馈和演化的完整闭环。它要解决的不是一次性把材料塞进提示词，而是把分散的数字行为持续转化为未来任务可直接复用的上下文数据产品。",
    "如图2-1所示，MetaFlow 首先从屏幕活动、键盘输入、鼠标操作、麦克风声音、文件变化、应用和终端事件等静默观察源持续收集数字环境中的行为变化。端侧只保留必要事实和来源信息，并优先在本地处理隐私敏感内容。随后系统把原始信号转化为“看到了什么、输入了什么、点击了什么、听到了什么、打开了什么、改变了什么”等结构化观察结果，降低后续模型理解的噪声。":
        "如图2-1所示，MetaFlow 首先从浏览器页面、屏幕活动、键盘输入、文件变化、项目状态、终端日志、AI 会话和用户反馈中生成 Observation。Observation 是事实入口，只回答“发生了什么、来自哪里、何时发生、如何采集”，不急于判断其长期意义。端侧优先保留来源、权限、时间和局部证据，并在本地处理隐私敏感内容，为后续 View 生成保留可追溯基础。",
    "在理解处理阶段，规则引擎、大模型、智能代理和命令行工具共同把观察结果加工为有意义的内部视图，例如当前页面、当前关注任务、当前项目、待办事项、建议、执行结果、每日记忆和长期偏好。内部视图不等同于简单摘要，而是面向未来任务组织的上下文坐标系。智能代理和用户界面都通过读取、更新或审核这些视图来采取行动，并把新的点击、编辑、拒绝、完成、失败等结果再次写回观察流，形成可持续学习的闭环。":
        "在 processor runtime 中，确定性代码、脚本、LLM prompt、agent task、browser job 和批处理任务共同把 Observation 与已有 Views 加工为新的 Views，例如 state.surface、work.focus_set、project.current、task.*、result.*、feedback.*、memory.daily 和 memory.profile。View 不是普通摘要，而是面向未来任务的坐标系：它应该减少下一次任务的搜索步骤、上下文拼接成本、模型输入量和失败重来次数。用户界面和 Agent 都读取同一套 ViewGraph，并把接受、修改、拒绝、完成、失败、复用等结果写回反馈观察。",
    "在工程实现上，MetaFlow 总体架构由端侧感知层、本地上下文层、内部视图层、云端智能层和反馈学习层组成。端侧负责浏览器页面、编辑器、终端、屏幕证据、本地记忆和上下文数据库；云端依托无问芯穹平台完成跨事件归纳、复杂候选生成、摘要压缩、任务规划和长程智能代理执行。端侧与云端通过视图图谱（ViewGraph）连接：端侧持续产生可追溯的观察结果和低延迟内部视图，云端在必要时对视图进行语义压缩、任务规划或演化候选生成，处理结果再写回视图图谱供后续复用。":
        "在工程实现上，MetaFlow 由 Observation stream、Task discovery、Processor runtime、ViewGraph、Personal applications、Verification 和 Evolution 七层组成。端侧负责浏览器页面、编辑器、终端、屏幕证据、本地记忆和上下文数据库；无问芯穹平台负责长上下文理解、多模态处理、批量语义压缩、智能代理探索和复杂候选生成。端侧与云端通过 ViewGraph 连接：端侧持续产生可追溯 Observation 和低延迟 Views，云端在必要时对任务证据包进行复杂推理、规划或压缩，结果再带着来源、运行事件和反馈写回 ViewGraph。",
}

storyline_note = """\n## 报告叙事主线建议\n\n一句话：MetaFlow 不是又一个会聊天或会点网页的 Agent，而是把长期工作痕迹转化为可检查、可复用、可验证 Views 的本地优先上下文运行时。\n\n推荐故事顺序：\n\n1. 行业变化：Agent 已经从回答问题走向使用浏览器、桌面和工具，但 benchmark 与商业产品都暴露出可靠性、状态和权限问题。\n2. 真实痛点：知识工作上下文散落在网页、文档、项目、日志、截图、对话和反馈里，每次 Agent 任务都像重新开局。\n3. 核心洞察：memory 不是历史仓库，而是能降低未来搜索成本的 ViewGraph；message 是发生过的事，View 是下一次工作能站上去的状态。\n4. 系统解法：Observation -> Processor -> ViewGraph -> 应用/Agent 行动 -> Feedback -> Evolution。\n5. 差异化：应用是 ViewGraph 的投影，Agent 是能力/worker，不是事实源；输出必须带 provenance、policy、trace 和 feedback。\n6. 赛题价值：端侧负责事实、权限、低延迟和本地证据；无问芯穹平台负责长上下文、多模态、智能代理探索和批量压缩。\n7. 评估逻辑：不是只看生成文本，而是看 View 是否让下一次任务更便宜：搜索步骤、token、时间、失败率、人工编辑和复用率是否下降。\n\n可反复使用的金句：\n\n- 让上下文可检查，让记忆可行动，让反馈推动系统变聪明。\n- Agent 是 transient，View 和 workflow 才是可复用资产。\n- 不可验证的 Agent 只能演示，可验证的 Agent 才能成为运行时。\n- 应用不是 memory 孤岛，而是同一个 ViewGraph 的不同投影。\n- 一个好的 View 的价值，不是保存了很多东西，而是减少下一次工作的搜索成本。\n"""


def main():
    if OUT.exists():
        os.chmod(OUT, 0o644)
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    os.chmod(OUT, 0o644)
    doc = Document(OUT)

    for old, new in replacements.items():
        replace_paragraph_text(doc, old, new)

    rename_heading(doc, "1.4 关键创新", "1.5 关键创新")
    rename_heading(doc, "1.5 与赛题要求的契合度", "1.7 与赛题要求的契合度")

    anchor = find_paragraph(doc, "1.5 关键创新")
    current = add_heading_before(anchor, aha_heading, level=2)
    for text in aha_paragraphs:
        current = add_paragraph_after(current, text)

    anchor = find_paragraph(doc, "1.7 与赛题要求的契合度")
    current = add_heading_before(anchor, related_heading, level=2)
    for text in related_intro:
        current = add_paragraph_after(current, text)
    current = add_paragraph_after(current, "表1-1 代表性开源项目与 MetaFlow 差异化定位", first_line=False)
    table = add_table_after(
        current,
        ["项目", "定位", "对 MetaFlow 的启发与差异", "引用"],
        open_source_rows,
        widths=[1.15, 2.0, 3.2, 0.65],
    )
    current = add_paragraph_after_table(table, "表1-2 代表性商业产品与 MetaFlow 差异化定位", first_line=False)
    add_table_after(
        current,
        ["产品", "定位", "MetaFlow 的相对优势", "引用"],
        commercial_rows,
        widths=[1.25, 2.15, 3.0, 0.6],
    )

    anchor = find_heading_startswith(doc, "5.3 技术效果评估")
    current = anchor
    for text in benchmark_paragraphs:
        current = add_paragraph_after(current, text)
    current = add_paragraph_after(current, "表5-2 公开 Agent 基准对 MetaFlow 测试方案的启发", first_line=False)
    add_table_after(
        current,
        ["基准", "任务环境", "对 MetaFlow 的启发", "引用"],
        [
            ("WebArena / Mind2Web", "真实网站、多领域网页任务", "纳入网页状态、弹窗、失败恢复、轨迹回放和来源审计。", "[24][25]"),
            ("OSWorld / AndroidWorld", "桌面与移动端真实应用", "将屏幕、应用状态、触控/GUI grounding 统一进入多模态 View。", "[28][29]"),
            ("GAIA / AgentBench", "工具使用、多轮决策、通用助手任务", "评估应覆盖中间过程、工具调用质量和人类反馈，而非只看最终文本。", "[27][30]"),
            ("SWE-bench / WorkArena", "真实代码 issue 与企业系统任务", "将可执行测试、业务规则和任务完成证据纳入工作流验收。", "[42][43]"),
            ("ToolSandbox / tau-bench", "状态化工具和用户交互", "把隐式状态、政策约束、澄清与中间里程碑作为可验证对象。", "[44][45]"),
        ],
        widths=[1.45, 1.8, 3.2, 0.55],
    )

    anchor = find_paragraph(doc, "6.2 商业化潜力")
    current = anchor
    for text in [
        "从竞品格局看，MetaFlow 的商业化切入点不应表述为“又一个通用 AI 助手”，而应表述为“个人与团队上下文智能运行层”。对于 B 端客户，价值来自可私有化部署、可审计证据链、可接入现有工具和可沉淀团队工作流；对于 C 端重度知识工作者，价值来自跨网页、文件、终端、对话和写作过程的连续上下文。与 Manus、Genspark、Replit Agent、Devin 等闭源托管产品相比，MetaFlow 更适合强调开源可控、端侧隐私、协议适配、可复现评测和竞赛/科研场景落地[38-41]。",
        "因此，商业化路径可以分为三阶段：第一阶段提供研究写作、项目续接、网页资料整理等本地优先的个人版；第二阶段面向研发、咨询、投研、法务、运营团队提供团队上下文中台和私有化部署；第三阶段将浏览器执行器、MCP/ACP adapter、Firecrawl 数据获取、云端模型路由和评测回归封装为可扩展 runtime，形成面向开发者和行业方案商的插件生态。",
    ]:
        current = add_paragraph_after(current, text)

    ref_heading = find_paragraph(doc, "参考文献")
    # Remove existing reference paragraphs after the heading.
    after_refs = False
    for p in list(doc.paragraphs):
        if p._p is ref_heading._p:
            after_refs = True
            continue
        if after_refs and p.text.strip().startswith("["):
            p._element.getparent().remove(p._element)
    for ref in references:
        add_reference_paragraph(doc, ref)

    # Lightly normalize added/renamed heading fonts across the new headings.
    for p in doc.paragraphs:
        if p.text.strip() in {aha_heading, related_heading, "1.5 关键创新", "1.7 与赛题要求的契合度"}:
            style_heading(p, level=2)

    doc.save(OUT)

    INDEX.write_text(
        "# MetaFlow Deep Research 引用与竞品索引\n\n"
        "## Aha Moment 主线\n\n"
        "- MetaFlow 不是又一个 Agent demo，而是 runtime/adapter layer：把脚本、浏览器、CLI、外部 Agent、个人应用和长期上下文统一到可组合、可验证、可跨端运行的系统中。\n"
        "- 标准边界：MCP/ACP/A2A 说明行业从单点 Agent 走向协议化互操作。\n"
        "- View memory：把上下文物化为可检查、可复用的 ViewGraph，而不是把历史简单塞进向量库。\n"
        "- 可验证运行：tracing、tool call lifecycle、权限、状态更新和人工接管是一等公民。\n"
        "- 可进化工作流：脚本提供确定性，Agent 提供探索性，ViewGraph 提供记忆，feedback 提供演化方向。\n\n"
        +
        storyline_note
        + "\n## 参考文献清单\n\n" + "\n".join(f"- {r}" for r in references) + "\n",
        encoding="utf-8",
    )

    print(OUT)
    print(INDEX)


if __name__ == "__main__":
    main()
